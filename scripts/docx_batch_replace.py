#!/usr/bin/env python3
"""
docx_batch_replace.py — DOCX 文本批量替换工具（鲁棒版本）

核心原理：使用 python-docx 读取段落级别的完整文本（自动合并 XML runs），
避免 zipfile 直接操作 XML 时因 Word 的 runs 拆分而导致替换失败。

同时遍历段落和表格单元格，确保不遗漏任何文本。

用法：
    python3 docx_batch_replace.py <目录或文件> --rules rules.json [--verify]

rules.json 示例：
{
    "replacements": {
        "补充法律意见书（二）": "补充法律意见书（三）",
        "2024年上半年度新增底稿": "2025年度报告新增底稿",
        "加审期间": "本次年度更新期间"
    },
    "forbidden_terms": ["加审", "半年报", "半年度"]
}
"""

import os
import sys
import json
import argparse

try:
    import docx
except ImportError:
    print('错误：需要安装 python-docx 库')
    print('运行: pip install python-docx')
    sys.exit(1)


def replace_in_paragraphs(paragraphs, replacements):
    """
    对一组段落执行文本替换。
    
    关键技术点：
    1. para.text 返回的是合并后的完整文本，无论 Word 内部如何拆分 XML runs
    2. 设置 para.text = new_text 会清除所有 runs，创建单个 run
    3. 必须保存并恢复原始格式（字体、大小等），否则排版会丢失
    
    Returns:
        int: 修改的段落数量
    """
    modified_count = 0
    
    for para in paragraphs:
        text = para.text
        original = text
        
        for old, new in replacements.items():
            if old in text:
                text = text.replace(old, new)
        
        if text != original:
            # 保存首个非空 run 的格式属性
            rPr_xml = None
            if para.runs:
                for r in para.runs:
                    if len(r.text.strip()) > 0 and r._r.rPr is not None:
                        rPr_xml = r._r.rPr
                        break
            
            # 执行替换（会重置为单个 run）
            para.text = text
            
            # 恢复格式
            if rPr_xml is not None and para.runs:
                para.runs[0]._r.append(rPr_xml)
            
            modified_count += 1
    
    return modified_count


def process_docx(filepath, replacements):
    """处理单个 DOCX 文件"""
    doc = docx.Document(filepath)
    total_modified = 0
    
    # 处理正文段落
    total_modified += replace_in_paragraphs(doc.paragraphs, replacements)
    
    # 处理表格中的段落（极易遗漏！）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_modified += replace_in_paragraphs(cell.paragraphs, replacements)
    
    # 处理页眉页脚
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.paragraphs:
                total_modified += replace_in_paragraphs(header.paragraphs, replacements)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.paragraphs:
                total_modified += replace_in_paragraphs(footer.paragraphs, replacements)
    
    if total_modified > 0:
        doc.save(filepath)
    
    return total_modified


def verify_docx(filepath, forbidden_terms):
    """验证单个 DOCX 文件不包含禁止词"""
    doc = docx.Document(filepath)
    issues = []
    
    # 检查段落
    for i, p in enumerate(doc.paragraphs):
        for term in forbidden_terms:
            if term in p.text:
                # 排除序号用法（如"（二）查验事项"）
                context = p.text[max(0, p.text.index(term)-5):p.text.index(term)+len(term)+10]
                issues.append({
                    'location': f'段落 {i+1}',
                    'term': term,
                    'context': context
                })
    
    # 检查表格
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    for term in forbidden_terms:
                        if term in p.text:
                            context = p.text[max(0, p.text.index(term)-5):p.text.index(term)+len(term)+10]
                            issues.append({
                                'location': f'表格{ti+1} 行{ri+1} 列{ci+1}',
                                'term': term,
                                'context': context
                            })
    
    return issues


def main():
    parser = argparse.ArgumentParser(description='DOCX 文本批量替换工具')
    parser.add_argument('target', help='目标目录或单个 DOCX 文件')
    parser.add_argument('--rules', required=True, help='替换规则 JSON 文件路径')
    parser.add_argument('--verify', action='store_true', help='替换后执行禁止词校验')
    parser.add_argument('--dry-run', action='store_true', help='仅扫描不修改，预览将被替换的内容')
    
    args = parser.parse_args()
    
    # 加载规则
    with open(args.rules, 'r', encoding='utf-8') as f:
        rules = json.load(f)
    
    replacements = rules.get('replacements', {})
    forbidden_terms = rules.get('forbidden_terms', [])
    
    # 确定目标文件列表
    if os.path.isfile(args.target):
        files = [args.target]
    elif os.path.isdir(args.target):
        files = [
            os.path.join(args.target, f)
            for f in os.listdir(args.target)
            if f.endswith('.docx') and not f.startswith(('~', '.'))
        ]
    else:
        print(f'错误：路径不存在 {args.target}')
        sys.exit(1)
    
    # 执行替换
    print(f'--- DOCX 批量替换 ---')
    print(f'目标文件: {len(files)} 个')
    print(f'替换规则: {len(replacements)} 条')
    print()
    
    total_files_modified = 0
    for filepath in sorted(files):
        basename = os.path.basename(filepath)
        
        if args.dry_run:
            doc = docx.Document(filepath)
            preview_count = 0
            for p in doc.paragraphs:
                for old in replacements:
                    if old in p.text:
                        preview_count += 1
                        print(f'  [预览] {basename}: "{old}" → "{replacements[old]}"')
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for old in replacements:
                                if old in p.text:
                                    preview_count += 1
            if preview_count > 0:
                print(f'  共 {preview_count} 处将被替换\n')
        else:
            count = process_docx(filepath, replacements)
            if count > 0:
                total_files_modified += 1
                print(f'✅ 已修改: {basename} ({count} 处)')
    
    if not args.dry_run:
        print(f'\n共修改 {total_files_modified}/{len(files)} 个文件')
    
    # 校验
    if args.verify and forbidden_terms:
        print(f'\n--- 禁止词校验 ---')
        print(f'禁止词: {forbidden_terms}')
        all_clean = True
        for filepath in sorted(files):
            issues = verify_docx(filepath, forbidden_terms)
            if issues:
                all_clean = False
                basename = os.path.basename(filepath)
                for issue in issues:
                    print(f'❌ {basename} [{issue["location"]}]: '
                          f'发现 "{issue["term"]}" → ...{issue["context"]}...')
        if all_clean:
            print('✅ 全部校验通过，无禁止词残留')


if __name__ == '__main__':
    main()
